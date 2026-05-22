"""
Smart Dentsu Buddy: Agentic RAG Chatbot for Marketing & Advertising Research
==============================================================================

A production-ready Streamlit chatbot powered by LangGraph and multi-source Agentic RAG.

Features:
- Multi-source knowledge retrieval (campaigns DB, research PDFs, web articles, live search)
- Marketing scope guardrails to prevent off-topic queries
- Self-correcting agent with query rewriting and relevance grading
- File upload support (PDF, images, Word docs)
- Persistent chat history with export
- Environment variables via Streamlit UI (no .env file)

Author: Dentsu AI Team
Last Updated: 2025-05-22
"""

import os
os.environ["PROTOCOL_BUFFERS_PYTHON_IMPLEMENTATION"] = "python"
import json
import time
import warnings
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Annotated, Sequence, TypedDict, Literal, Any

import streamlit as st
from streamlit_option_menu import option_menu

# Document loading and processing
from langchain_community.document_loaders import UnstructuredFileLoader, WebBaseLoader, PyMuPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# Embeddings and vector store
from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
from langchain_chroma import Chroma

# Prompt templates and output parsers
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.messages import BaseMessage, HumanMessage, SystemMessage
from langchain_core.tools import tool, create_retriever_tool

# LangGraph components
from langgraph.graph import StateGraph, START, END
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode, tools_condition

# Structured output
from pydantic import BaseModel, Field

# Web search
from langchain_community.tools.tavily_search import TavilySearchResults

# Ignore warnings
warnings.filterwarnings("ignore")

# ============================================================================
# STREAMLIT PAGE CONFIGURATION
# ============================================================================

st.set_page_config(
    page_title="Smart Dentsu Buddy",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Custom CSS for better UI
st.markdown(
    """
    <style>
    .main {
        padding: 20px;
    }
    .header-gradient {
        background: linear-gradient(135deg, #000000 0%, #1A1A2E 50%, #E30613 100%);
        padding: 30px;
        border-radius: 12px;
        color: white;
        text-align: center;
        margin-bottom: 30px;
    }
    .chat-message {
        padding: 12px 16px;
        margin: 8px 0;
        border-radius: 8px;
    }
    .user-message {
        background-color: #E8F0FE;
        border-left: 4px solid #1f77e1;
    }
    .assistant-message {
        background-color: #F0F5F9;
        border-left: 4px solid #34A853;
    }
    .system-message {
        background-color: #FEF5E7;
        border-left: 4px solid #F39C12;
    }
    .error-message {
        background-color: #FADBD8;
        border-left: 4px solid #E74C3C;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ============================================================================
# SESSION STATE INITIALIZATION
# ============================================================================

def initialize_session_state():
    """Initialize all required session state variables."""
    if "credentials_set" not in st.session_state:
        st.session_state.credentials_set = False
    if "llm" not in st.session_state:
        st.session_state.llm = None
    if "embeddings" not in st.session_state:
        st.session_state.embeddings = None
    if "graph" not in st.session_state:
        st.session_state.graph = None
    if "tools" not in st.session_state:
        st.session_state.tools = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "conversation_id" not in st.session_state:
        st.session_state.conversation_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    if "uploaded_files" not in st.session_state:
        st.session_state.uploaded_files = []

initialize_session_state()

# ============================================================================
# CONFIGURATION & CREDENTIALS (SIDEBAR)
# ============================================================================

def render_sidebar():
    """Render sidebar with configuration, credentials, and controls."""
    with st.sidebar:
        st.markdown("## ⚙️ Configuration")
        
        # Tabs for configuration
        config_tab, history_tab = st.tabs(["API Keys", "History"])
        
        with config_tab:
            st.markdown("### Azure OpenAI Setup")
            
            # Azure OpenAI credentials
            endpoint = st.text_input(
                "Azure OpenAI Endpoint",
                placeholder="https://<resource>.openai.azure.com/",
                type="password",
                key="endpoint_input",
            )
            
            api_key = st.text_input(
                "Azure OpenAI API Key",
                type="password",
                key="api_key_input",
            )
            
            chat_model = st.text_input(
                "Chat Model Name",
                value="gpt-4o",
                placeholder="e.g., gpt-4o",
                key="chat_model_input",
            )
            
            embedding_model = st.text_input(
                "Embedding Model Name",
                value="text-embedding-ada-002",
                placeholder="e.g., text-embedding-ada-002",
                key="embedding_model_input",
            )
            
            api_version = st.text_input(
                "API Version",
                value="2024-02-15-preview",
                placeholder="e.g., 2024-02-15-preview",
                key="api_version_input",
            )
            
            st.markdown("### Tavily (Web Search)")
            
            tavily_key = st.text_input(
                "Tavily API Key",
                type="password",
                placeholder="Optional - for live web search",
                key="tavily_key_input",
            )
            
            # Initialize button
            if st.button("🚀 Initialize Agent", use_container_width=True):
                if not endpoint or not api_key or not chat_model or not embedding_model:
                    st.error("❌ Please fill in all required Azure OpenAI fields")
                else:
                    with st.spinner("Initializing agent and loading knowledge bases..."):
                        try:
                            setup_agent(
                                endpoint=endpoint,
                                api_key=api_key,
                                chat_model=chat_model,
                                embedding_model=embedding_model,
                                api_version=api_version,
                                tavily_key=tavily_key,
                            )
                            st.success("✅ Agent initialized successfully!")
                            st.session_state.credentials_set = True
                        except Exception as e:
                            st.error(f"❌ Error: {str(e)}")
        
        with history_tab:
            st.markdown("### Conversation History")
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("📝 New Conversation", use_container_width=True):
                    st.session_state.chat_history = []
                    st.session_state.conversation_id = datetime.now().strftime("%Y%m%d_%H%M%S")
                    st.success("Started new conversation")
            
            with col2:
                if st.button("💾 Save History", use_container_width=True):
                    save_chat_history()
                    st.success("Conversation saved to `chat_history.json`")
            
            # Show conversation stats
            if st.session_state.chat_history:
                st.markdown(f"**Messages:** {len(st.session_state.chat_history)}")
                st.markdown(f"**ID:** `{st.session_state.conversation_id}`")
        
        # File upload section
        st.markdown("### 📤 Upload Documents")
        uploaded_files = st.file_uploader(
            "Upload PDF, Image, or Word document",
            type=["pdf", "png", "jpg", "jpeg", "doc", "docx"],
            accept_multiple_files=True,
        )
        
        if uploaded_files:
            if st.button("Process Uploads", use_container_width=True):
                process_uploaded_files(uploaded_files)

# ============================================================================
# AGENT SETUP & INITIALIZATION
# ============================================================================

class MarketingDecision(BaseModel):
    """Structured output for guardrail classification."""
    decision: Literal["YES", "NO"] = Field(
        description="YES if marketing/advertising-related, NO otherwise"
    )

class AgentState(TypedDict):
    """Shared state passed between graph nodes."""
    messages: Annotated[Sequence[BaseMessage], add_messages]

def setup_agent(
    endpoint: str,
    api_key: str,
    chat_model: str,
    embedding_model: str,
    api_version: str,
    tavily_key: str = None,
):
    """Initialize LLM, embeddings, vector stores, tools, and graph."""
    # Create LLM and embeddings clients
    llm = AzureChatOpenAI(
        azure_deployment=chat_model,
        api_version=api_version,
        azure_endpoint=endpoint,
        api_key=api_key,
        temperature=0,
    )
    
    embeddings = AzureOpenAIEmbeddings(
        azure_endpoint=endpoint,
        azure_deployment=embedding_model,
        openai_api_version=api_version,
        api_key=api_key,
    )
    
    st.session_state.llm = llm
    st.session_state.embeddings = embeddings
    
    # ====================================================================
    # KNOWLEDGE BASE 1: CAMPAIGN DATABASE (JSON)
    # ====================================================================
    isExistCampaigns = os.path.exists("campaigns_db.json")
    st.markdown(f"**isExistCampaigns:** `{isExistCampaigns}`")
    if os.path.exists("campaigns_db.json"):
        with open("campaigns_db.json", "r") as f:
            campaigns_db_docs = json.load(f)
        
        campaign_documents = []
        for idx, campaign in enumerate(campaigns_db_docs):
            fields = "\n".join(f"{key}: {value}" for key, value in campaign.items())
            campaign_documents.append(
                Document(
                    page_content=fields,
                    metadata={
                        "source": "campaigns_db.json",
                        "index": idx,
                        "industry": campaign.get("industry", ""),
                        "client": campaign.get("client", ""),
                        "campaign_name": campaign.get("campaign_name", ""),
                    },
                )
            )
        
        PERSIST_DIR_CAMPAIGNS = "campaigns_db_vectorstore"
        if os.path.exists(PERSIST_DIR_CAMPAIGNS):
            vectordb_campaigns = Chroma(
                collection_name="campaigns-db",
                persist_directory=PERSIST_DIR_CAMPAIGNS,
                embedding_function=embeddings,
            )
        else:
            vectordb_campaigns = Chroma.from_documents(
                documents=campaign_documents,
                collection_name="campaigns-db",
                embedding=embeddings,
                persist_directory=PERSIST_DIR_CAMPAIGNS,
            )
        
        retriever_campaigns = vectordb_campaigns.as_retriever(
            search_type="similarity_score_threshold",
            search_kwargs={"score_threshold": 0.5, "k": 5},
        )
        
        retriever_tool_campaigns = create_retriever_tool(
            retriever=retriever_campaigns,
            name="search_campaigns_db",
            description="Search marketing campaigns, advertising strategies, and media plans. "
            "Use for campaign recommendations, performance benchmarks, and budget allocation examples.",
        )
        campaign_tool = retriever_tool_campaigns
    else:
        campaign_tool = None
    
    # ====================================================================
    # KNOWLEDGE BASE 2: RESEARCH PDFS
    # ====================================================================
    
    pdf_tool = None
    pdf_files = [
        "Content_Effects_Advertising_Marketing.pdf",
        "Digital_Transformation_in_Marketing.pdf",
    ]
    
    if all(os.path.exists(pdf) for pdf in pdf_files):
        all_pdf_docs = []
        for pdf_file in pdf_files:
            docs = PyMuPDFLoader(pdf_file).load_and_split()
            all_pdf_docs.extend(docs)
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=600, chunk_overlap=100
        )
        pdf_texts = text_splitter.split_documents(all_pdf_docs)
        
        PERSIST_DIR_PDF = "marketing_pdf_db"
        if os.path.exists(PERSIST_DIR_PDF):
            vectordb_pdf = Chroma(
                collection_name="marketing-pdf-docs",
                persist_directory=PERSIST_DIR_PDF,
                embedding_function=embeddings,
            )
        else:
            vectordb_pdf = Chroma.from_documents(
                documents=pdf_texts,
                collection_name="marketing-pdf-docs",
                embedding=embeddings,
                persist_directory=PERSIST_DIR_PDF,
            )
        
        retriever_pdf = vectordb_pdf.as_retriever(
            search_type="similarity_score_threshold",
            search_kwargs={"score_threshold": 0.5, "k": 5},
        )
        
        retriever_tool_pdf = create_retriever_tool(
            retriever=retriever_pdf,
            name="search_marketing_research",
            description="Search marketing research papers for content effectiveness, "
            "digital transformation, programmatic advertising, and marketing automation insights.",
        )
        pdf_tool = retriever_tool_pdf
    
    # ====================================================================
    # KNOWLEDGE BASE 3: MARKETING LAW ARTICLES
    # ====================================================================
    
    web_tool_static = None
    marketing_urls = [
        "https://tenthings.blog/2023/06/30/ten-things-marketing-law-basics-for-in-house-counsel/",
        "https://blog.ipleaders.in/marketing-media-consumer-protection-law-india/",
    ]
    
    try:
        docs = [WebBaseLoader(url).load() for url in marketing_urls]
        docs_list = [item for sublist in docs for item in sublist]
        
        if docs_list:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=600, chunk_overlap=100
            )
            web_texts = text_splitter.split_documents(docs_list)
            
            PERSIST_DIR_WEB = "marketing_law_articles_db"
            if os.path.exists(PERSIST_DIR_WEB):
                vectordb_web = Chroma(
                    collection_name="marketing-web-docs",
                    persist_directory=PERSIST_DIR_WEB,
                    embedding_function=embeddings,
                )
            else:
                vectordb_web = Chroma.from_documents(
                    documents=web_texts,
                    collection_name="marketing-web-docs",
                    embedding=embeddings,
                    persist_directory=PERSIST_DIR_WEB,
                )
            
            retriever_web = vectordb_web.as_retriever(
                search_type="similarity_score_threshold",
                search_kwargs={"score_threshold": 0.5, "k": 5},
            )
            
            retriever_tool_web = create_retriever_tool(
                retriever=retriever_web,
                name="search_marketing_law_articles",
                description="Search legal and compliance insights on marketing law, "
                "advertising compliance, endorsements, consumer protection in India.",
            )
            web_tool_static = retriever_tool_web
    except Exception as e:
        st.warning(f"⚠️ Could not load marketing law articles: {str(e)}")
    
    # ====================================================================
    # TOOL 4: LIVE WEB SEARCH (TAVILY)
    # ====================================================================
    
    search_web_tool = None
    st.markdown(f"**tavily_key Found:** `{tavily_key}`")
    if tavily_key:
        tavily_search = TavilySearchResults(
            tavily_api_key= tavily_key,
            max_results=5,
            search_depth="advanced",
            include_raw_content=True,
        )
        
        @tool
        def search_web(query: str) -> str:
            """Search the web for current information, news, and recent events."""
            results = tavily_search.invoke(query)
            formatted_results = []
            for r in results:
                title = r.get("title", "No Title")
                content = r.get("content", "No Content")
                url = r.get("url", "")
                formatted_results.append(f"Title: {title}\nContent: {content}\nSource: {url}")
            return "\n\n---\n\n".join(formatted_results) if formatted_results else "No results found."
        
        search_web_tool = search_web
    
    # ====================================================================
    # BUILD TOOLS LIST
    # ====================================================================
    
    tools = []
    if campaign_tool:
        tools.append(campaign_tool)
    if pdf_tool:
        tools.append(pdf_tool)
    if web_tool_static:
        tools.append(web_tool_static)
    if search_web_tool:
        tools.append(search_web_tool)
    
    st.session_state.tools = tools
    
    # ====================================================================
    # BUILD GRAPH NODES
    # ====================================================================
    
    def agent(state: AgentState):
        """Agent node: decides which tool to use."""
        model_with_tools = llm.bind_tools(tools) if tools else llm
        response = model_with_tools.invoke(state["messages"])
        return {"messages": [response]}
    
    def generate(state: AgentState):
        """Generate node: writes final answer."""
        messages = state["messages"]
        question = messages[0].content
        docs = messages[-1].content
        
        prompt = ChatPromptTemplate.from_messages([
            ("system",
             "You are a helpful marketing and advertising research assistant for Dentsu. "
             "Use the following documents to answer the question.\n\n"
             "IMPORTANT GUIDELINES:\n"
             "- Provide accurate, evidence-based information from the retrieved documents.\n"
             "- Include a disclaimer: 'This information is for research and strategic planning purposes only "
             "and should be validated with current market data before making investment decisions.'\n"
             "- Do NOT fabricate campaign metrics or performance data.\n"
             "- If someone asks about a live campaign, advise them to check the relevant platform dashboards for real-time data.\n"
             "- If the documents don't contain the answer, say so honestly.\n"
             "- Keep your answer clear, concise, and well-structured."),
            ("human",
             "Documents:\n{context}\n\n"
             "Question: {question}\n\n"
             "Answer:"),
        ])
        
        chain = prompt | llm | StrOutputParser()
        answer = chain.invoke({"context": docs, "question": question})
        return {"messages": [answer]}
    
    def rewrite(state: AgentState):
        """Rewrite node: improves question if retrieval fails."""
        original_question = state["messages"][0].content
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", "You are an expert at refining search queries. Rewrite the question to be "
             "clearer, more specific, and more likely to retrieve relevant documents."),
            ("human",
             "Original question: {question}\n\n"
             "Rewrite this as a clearer, more specific, and more searchable question:"),
        ])
        
        chain = prompt | llm
        response = chain.invoke({"question": original_question})
        return {"messages": [HumanMessage(content=response.content)]}
    
    def grade_documents(state: AgentState) -> Literal["generate", "rewrite"]:
        """Grade node: checks if retrieved docs are relevant."""
        class Grade(BaseModel):
            score: str = Field(description="'yes' if relevant, 'no' if not")
        
        grader = llm.with_structured_output(Grade)
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", "You are a document relevance grader. Determine if the retrieved document "
             "helps answer the user's question. Reply with 'yes' or 'no' only."),
            ("human",
             "Retrieved Document:\n{context}\n\n"
             "User Question: {question}\n\n"
             "Is this document relevant to the question? (yes/no)"),
        ])
        
        chain = prompt | grader
        question = state["messages"][0].content
        docs = state["messages"][-1].content
        
        result = chain.invoke({"question": question, "context": docs})
        return "generate" if result.score == "yes" else "rewrite"
    
    # Build graph
    workflow = StateGraph(AgentState)
    workflow.add_node("agent", agent)
    workflow.add_node("retrieve", ToolNode(tools) if tools else lambda x: x)
    workflow.add_node("generate", generate)
    workflow.add_node("rewrite", rewrite)
    
    workflow.add_edge(START, "agent")
    workflow.add_conditional_edges(
        "agent",
        tools_condition if tools else lambda x: END,
        {"tools": "retrieve", END: END} if tools else {END: END},
    )
    workflow.add_conditional_edges("retrieve", grade_documents)
    workflow.add_edge("generate", END)
    workflow.add_edge("rewrite", "agent")
    
    st.session_state.graph = workflow.compile()

# ============================================================================
# GUARDRAIL IMPLEMENTATION
# ============================================================================

def check_marketing_guardrail(question: str) -> dict:
    """Check if question is marketing/advertising-related."""
    if not st.session_state.llm:
        return {"approved": False, "message": "LLM not initialized"}
    
    guardrail_prompt = (
        "Classify if the following user query is related to marketing, advertising, "
        "media planning, brand strategy, campaign management, digital marketing, content strategy, "
        "programmatic advertising, media buying, creative strategy, audience targeting, marketing analytics, "
        "ROI measurement, social media marketing, SEO/SEM, influencer marketing, MarTech, CRM, "
        "advertising technology, marketing operations, or legal/compliance topics related to marketing/advertising. "
        "Return YES or NO only."
    )
    
    messages = [
        SystemMessage(content=guardrail_prompt),
        HumanMessage(content=question),
    ]
    
    llm_classifier = st.session_state.llm.with_structured_output(MarketingDecision)
    classifier_response = llm_classifier.invoke(messages)
    
    if classifier_response.decision != "YES":
        return {
            "approved": False,
            "message": (
                "**Out of Scope** ❌\n\n"
                "I can only assist with marketing, advertising, and related industry queries. "
                "Please try again with a marketing or advertising-related question."
            ),
        }
    
    return {"approved": True, "message": None}

# ============================================================================
# FILE UPLOAD & PROCESSING
# ============================================================================

def process_uploaded_files(uploaded_files):
    """Process and embed uploaded documents."""
    if not st.session_state.embeddings:
        st.error("❌ Please initialize the agent first")
        return
    
    progress_bar = st.progress(0)
    total_files = len(uploaded_files)
    
    for idx, uploaded_file in enumerate(uploaded_files):
        try:
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(uploaded_file.name).suffix) as tmp:
                tmp.write(uploaded_file.getbuffer())
                tmp_path = tmp.name
            
            # Load documents
            if uploaded_file.type == "application/pdf":
                docs = PyMuPDFLoader(tmp_path).load_and_split()
            elif uploaded_file.type.startswith("image/"):
                # For images, create a simple document with filename
                docs = [Document(page_content=f"Image: {uploaded_file.name}", metadata={"source": uploaded_file.name})]
            elif uploaded_file.type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                from docx import Document as DocxDocument
                doc = DocxDocument(tmp_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                docs = [Document(page_content=text, metadata={"source": uploaded_file.name})]
            else:
                st.warning(f"Unsupported file type: {uploaded_file.name}")
                continue
            
            # Embed documents
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=100)
            split_docs = text_splitter.split_documents(docs)
            
            PERSIST_DIR_UPLOADS = "uploaded_docs_vectorstore"
            if os.path.exists(PERSIST_DIR_UPLOADS):
                vectordb_uploads = Chroma(
                    collection_name="uploaded-docs",
                    persist_directory=PERSIST_DIR_UPLOADS,
                    embedding_function=st.session_state.embeddings,
                )
                vectordb_uploads.add_documents(split_docs)
            else:
                Chroma.from_documents(
                    documents=split_docs,
                    collection_name="uploaded-docs",
                    embedding=st.session_state.embeddings,
                    persist_directory=PERSIST_DIR_UPLOADS,
                )
            
            st.session_state.uploaded_files.append(uploaded_file.name)
            st.success(f"✅ Processed: {uploaded_file.name}")
            
            # Cleanup
            os.unlink(tmp_path)
        
        except Exception as e:
            st.error(f"❌ Error processing {uploaded_file.name}: {str(e)}")
        
        progress_bar.progress((idx + 1) / total_files)

# ============================================================================
# CHAT HISTORY MANAGEMENT
# ============================================================================

def save_chat_history():
    """Save chat history to JSON file."""
    history_data = {
        "conversation_id": st.session_state.conversation_id,
        "timestamp": datetime.now().isoformat(),
        "messages": [
            {
                "role": msg.get("role", "unknown"),
                "content": msg.get("content", ""),
                "timestamp": msg.get("timestamp", ""),
            }
            for msg in st.session_state.chat_history
        ],
    }
    
    with open("chat_history.json", "w") as f:
        json.dump(history_data, f, indent=2)

def load_chat_history():
    """Load chat history from JSON file."""
    if os.path.exists("chat_history.json"):
        with open("chat_history.json", "r") as f:
            return json.load(f)
    return None

# ============================================================================
# MAIN CHAT INTERFACE
# ============================================================================

def render_chat_interface():
    """Render the main chat interface."""
    
    # Header
    st.markdown(
        """
        <div class="header-gradient">
        <h1>🤖 Smart Dentsu Buddy</h1>
        <p>Agentic RAG for Marketing & Advertising Research</p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    
    # Check if agent is initialized
    if not st.session_state.credentials_set:
        st.warning(
            "⚠️ Please configure your API keys in the **Configuration** tab on the left sidebar "
            "and click **Initialize Agent** to get started."
        )
        return
    
    if not st.session_state.graph:
        st.error("❌ Agent not initialized. Please check your API credentials.")
        return
    
    # Display chat history
    chat_container = st.container()
    with chat_container:
        for msg in st.session_state.chat_history:
            if msg["role"] == "user":
                st.markdown(
                    f'<div class="chat-message user-message"><strong>You:</strong> {msg["content"]}</div>',
                    unsafe_allow_html=True,
                )
            elif msg["role"] == "assistant":
                st.markdown(
                    f'<div class="chat-message assistant-message"><strong>Buddy:</strong> {msg["content"]}</div>',
                    unsafe_allow_html=True,
                )
            elif msg["role"] == "system":
                st.markdown(
                    f'<div class="chat-message system-message"><strong>System:</strong> {msg["content"]}</div>',
                    unsafe_allow_html=True,
                )
    
    # Input form
    st.markdown("---")
    
    col1, col2 = st.columns([0.9, 0.1])
    with col1:
        user_input = st.text_input(
            "Ask a question about marketing, advertising, or campaigns...",
            placeholder="e.g., Find me a successful digital campaign for CPG...",
            key="user_input",
        )
    with col2:
        submit_button = st.button("Send ➤", use_container_width=True)
    
    # Process user input
    if submit_button and user_input:
        # Add user message to history
        st.session_state.chat_history.append({
            "role": "user",
            "content": user_input,
            "timestamp": datetime.now().isoformat(),
        })
        
        # Check guardrail
        guardrail_result = check_marketing_guardrail(user_input)
        
        if not guardrail_result["approved"]:
            st.session_state.chat_history.append({
                "role": "system",
                "content": guardrail_result["message"],
                "timestamp": datetime.now().isoformat(),
            })
            st.warning(guardrail_result["message"])
            save_chat_history()
            st.rerun()
        
        # Run agent
        with st.spinner("🔍 Searching knowledge bases and reasoning..."):
            try:
                inputs = {"messages": [("user", user_input)]}
                result = st.session_state.graph.invoke(inputs)
                
                # Extract answer
                answer = result["messages"][-1].content if result["messages"] else "No response generated"
                st.markdown("Answer:", answer)

                # Add assistant response to history
                st.session_state.chat_history.append({
                    "role": "assistant",
                    "content": answer,
                    "timestamp": datetime.now().isoformat(),
                })
                
                # Save history
                save_chat_history()
                
                # Rerun to display new messages
                st.rerun()
            
            except Exception as e:
                error_msg = f"Error during query processing: {str(e)}"
                st.session_state.chat_history.append({
                    "role": "system",
                    "content": error_msg,
                    "timestamp": datetime.now().isoformat(),
                })
                st.error(f"❌ {error_msg}")
                save_chat_history()

# ============================================================================
# MAIN APP ENTRY POINT
# ============================================================================

def main():
    """Main app entry point."""
    render_sidebar()
    render_chat_interface()

if __name__ == "__main__":
    main()
